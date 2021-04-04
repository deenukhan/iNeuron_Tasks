import os
import shutil
from docx import Document

work_path = os.getcwd() # let's set the working path

#Let's Create the two folders
os.mkdir('theoritical_assignments')
os.mkdir('programming_assignments')

# Let's copy files from "python_basic" to theoritical_assignments

python_basic_files = os.listdir(os.path.join(work_path, 'python_basic'))

for file_name in python_basic_files:
    #Here we're doing nothing fancy just copy files from one location to another
    source = os.path.join(os.path.join(work_path, 'python_basic'), file_name)
    destination = os.path.join(work_path, 'theoritical_assignments')
    shutil.copy(source, destination)
    
# Let's copy files from "python_advance", to theoritical_assignments and parallely renaming them\
# We have already copied 25 assignments into this folder, now we will rename each file starting from 26th
# Eg. Assignment_1.docx in python_advance folder will become Assignment_26.docx in theoritical_assignments folder

python_advance_files = os.listdir(os.path.join(work_path, 'python_advance'))

for file_name in python_advance_files:
    #Below Two Lines are for just extracting the file number so that we can rename the file after 25 and so on.
    file_number = file_name.split('.')[0].split('_')[1]
    new_file_name = "Assignment_"+ str(int(file_number) + 25) + ".docx"
    source = os.path.join(os.path.join(work_path, 'python_advance'), file_name)
    destination = os.path.join(os.path.join(work_path, 'theoritical_assignments'), new_file_name)
    shutil.move(source, destination)

# Let's copy files from "python_programming_basic",  to programming_assignments
python_basic_program_files = os.listdir(os.path.join(work_path, 'python_programing_basic'))

for file_name in python_basic_program_files:
    
    #Here we're doing nothing fancy just copy files from one location to another
    source = os.path.join(os.path.join(work_path, 'python_programing_basic'), file_name)
    destination = os.path.join(work_path, 'programming_assignments')
    shutil.copy(source, destination)

# Let's copy files from "python_programming_advance",  to programming_assignments
# Here we don't need to rename the files as they are already different
python_adv_program_files = os.listdir(os.path.join(work_path, 'python_programming_advance'))

for file_name in python_adv_program_files:
    #Here we're doing nothing fancy just copy files from one location to another
    source = os.path.join(os.path.join(work_path, 'python_programming_advance'), file_name)
    destination = os.path.join(work_path, 'programming_assignments')
    shutil.copy(source, destination)


# Here we will be using external library docx's Document() Class deal with docx files
# Below code is just for reading the document file and copying it into another new document file
os.chdir(os.path.join(work_path, 'theoritical_assignments'))
theory_docx_files = os.listdir()
new_file = Document()           #Creating the object for new docx file

#Here We are just looping through each docx file and reading it's content one by one copying it to another file
for file in theory_docx_files:
    
    #Below Line is just for typing the heading
    new_file.add_heading(file, level = 0)
    doc = Document(file)
    paragraphs = doc.paragraphs
    #Here we will not be copying the blank lines, because i saw many docx are having multiple blank lines in between
    for para in paragraphs:
        if len(para.text) == 0:
            continue
        else:
            new_file.add_paragraph(para.text)

# Just Saving the file
new_file.save(os.path.join(work_path, "combined_theory_assignments.docx"))

# Here we will be using external library docx's Document() Class deal with docx files
# Below code is just for reading the document file and copying it into another new document file
os.chdir(os.path.join(work_path, 'programming_assignments'))
programming_docx_files = os.listdir()
new_file = Document()                #Creating the object for new docx file

#Here We are just looping through each docx file and reading it's content one by one copying it to another file
for file in programming_docx_files:
    
    #Below Line is just for typing the heading
    new_file.add_heading(file, level = 0)
    doc = Document(file)
    paragraphs = doc.paragraphs
    #Here we will not be copying the blank lines, because i saw many docx are having multiple blank lines in between
    for para in paragraphs:
        if len(para.text) == 0:
            continue
        else:
            new_file.add_paragraph(para.text)
            
# Just Saving the file
new_file.save(os.path.join(work_path, "combined_programming_assignments.docx"))

#Now We will be deleting all the folders except
folders_list = ['python_advance', 'python_basic', 'python_programing_basic', 'python_programming_advance', 'theoritical_assignments', 'programming_assignments']
os.chdir(work_path)
for folders in folders_list:
    shutil.rmtree(folders)